from fastapi import APIRouter, Depends, HTTPException, Form, UploadFile, File
from sqlalchemy.orm import Session
from db.models import RoleEnum, LicenseEnum, LanguageTranslation, User
from db.database import get_db
from auth.auth_manager import AuthManager
from fastapi.responses import FileResponse, PlainTextResponse
from transformers import pipeline
from docx import Document
import tempfile
import fitz
import docx
import easyocr
import numpy as np
import cv2
import io
from PIL import Image

router = APIRouter()

@router.post("/usecase/language-translation", response_model=None)
@AuthManager.check_access([RoleEnum.Editor], [LicenseEnum.Enterprise])
async def language_translation(
    input_lang: str = Form(...),
    output_lang: str = Form(...),
    text_input: str = Form(None),
    file: UploadFile = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(AuthManager.get_current_user),
    download_filetype: str = Form(None)
):
    has_text = text_input is not None and text_input.strip() != ""
    has_file = file is not None

    # Ensure only one input is provided
    if (has_text and has_file) or (not has_text and not has_file):
        raise HTTPException(
            status_code=400,
            detail="Please provide either text input or a file, but not both."
        )

    text = text_input if has_text else ""
    if has_file:
        filename = file.filename.lower()
        contents = await file.read()
        if not contents:
            raise HTTPException(status_code=400, detail="Uploaded file is empty.")
        if filename.endswith(".txt"):
            text = contents.decode("utf-8")
        elif filename.endswith(".pdf"):
            doc = fitz.open(stream=contents, filetype="pdf")
            reader = easyocr.Reader(['en'], gpu=False)
            for page in doc:
                text += page.get_text()
                images = page.get_images(full=True)
                for img_index, img in enumerate(images):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image = Image.open(io.BytesIO(image_bytes)).convert('RGB')
                    img_np = np.array(image)
                    img_cv = cv2.cvtColor(img_np, cv2.COLOR_RGB2BGR)
                    ocr_result = reader.readtext(img_cv, detail=0)
                    if ocr_result:
                        text += " " + " ".join(ocr_result)
        elif filename.endswith(".docx"):
            docx_doc = docx.Document(io.BytesIO(contents))
            reader = easyocr.Reader(['en'], gpu=False)
            for para in docx_doc.paragraphs:
                text += para.text + " "
            for rel in docx_doc.part._rels:
                rel_obj = docx_doc.part._rels[rel]
                if "image" in rel_obj.target_ref:
                    image_data = rel_obj.target_part.blob
                    image = Image.open(io.BytesIO(image_data)).convert('RGB')
                    img_np = np.array(image)
                    img_cv = cv2.cvtColor(img_np, cv2.COLOR_RGB2BGR)
                    ocr_result = reader.readtext(img_cv, detail=0)
                    if ocr_result:
                        text += " " + " ".join(ocr_result)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Use TXT, PDF, or DOCX.")

    # LOGGING: Print the extracted text for debugging
    print("[DEBUG] Extracted text from file:", repr(text))
    text = text.strip()
    if not text:
        print("[DEBUG] No text found after extraction.")
        raise HTTPException(status_code=400, detail="No text found in the document or input.")
    match = db.query(LanguageTranslation).filter(
        LanguageTranslation.input_lang == input_lang.lower(),
        LanguageTranslation.output_lang == output_lang.lower(),
        LanguageTranslation.input_text.ilike(f"%{text.strip()}%")
    ).first()
    if match:
        translated_text = match.output_text
        if download_filetype == "docx":
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                doc = Document()
                safe_text = translated_text if translated_text else ""
                doc.add_paragraph(safe_text)
                doc.save(tmp.name)
                tmp.flush()
                return FileResponse(
                    path=tmp.name,
                    filename="translated.docx",
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        elif download_filetype in ("txt", "txt-download"):
            with tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8') as tmp:
                safe_text = translated_text if translated_text else ""
                tmp.write(safe_text)
                tmp.flush()
                return FileResponse(
                    path=tmp.name,
                    filename="translated.txt",
                    media_type="text/plain; charset=utf-8"
                )
        return {"translated_text": translated_text}
    else:
        # Fallback: Use HuggingFace transformers pipeline for translation
        input_lang_code = input_lang.lower()
        output_lang_code = output_lang.lower()
        try:
            model_name = f"Helsinki-NLP/opus-mt-{input_lang_code}-{output_lang_code}"
            translator = pipeline(
                "translation",
                model=model_name
            )
            # Split text into chunks of max 800 characters (to stay under token limit)
            def split_text(text, max_chars=800):
                chunks = []
                for i in range(0, len(text), max_chars):
                    chunks.append(text[i:i+max_chars])
                return chunks
            text_chunks = split_text(text)
            translated_chunks = []
            for chunk in text_chunks:
                translated = translator(chunk, max_length=512)
                translated_text = translated[0]['translation_text'] if translated and 'translation_text' in translated[0] else ""
                translated_chunks.append(translated_text)
            translated_text = '\n'.join(translated_chunks)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

        if download_filetype == "docx":
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                doc = Document()
                safe_text = translated_text if translated_text else ""
                doc.add_paragraph(safe_text)
                doc.save(tmp.name)
                tmp.flush()
                return FileResponse(
                    path=tmp.name,
                    filename="translated.docx",
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        elif download_filetype == "txt":
            with tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8') as tmp:
                safe_text = translated_text if translated_text else ""
                tmp.write(safe_text)
                tmp.flush()
                return FileResponse(
                    path=tmp.name,
                    filename="translated.txt",
                    media_type="text/plain; charset=utf-8"
                )
        elif download_filetype == "txt-download":
            with tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8') as tmp:
                safe_text = translated_text if translated_text else ""
                tmp.write(safe_text)
                tmp.flush()
                return FileResponse(
                    path=tmp.name,
                    filename="translated.txt",
                    media_type="text/plain; charset=utf-8"
                )
        return {"translated_text": translated_text}
